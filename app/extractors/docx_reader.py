from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models import TextBlock


def _run_bold(paragraph) -> bool:  # noqa: ANN001
    if not paragraph.runs:
        return False
    bold_chars = 0
    total = 0
    for run in paragraph.runs:
        text = run.text or ""
        if not text.strip():
            continue
        total += len(text)
        if run.bold:
            bold_chars += len(text)
    if total == 0:
        return False
    return bold_chars / total >= 0.6


def _run_font_size(paragraph) -> float | None:  # noqa: ANN001
    sizes: list[float] = []
    for run in paragraph.runs:
        if run.font and run.font.size is not None:
            sizes.append(run.font.size.pt)
    if not sizes:
        # Tenta estilo do parágrafo
        try:
            style_size = paragraph.style.font.size
            if style_size is not None:
                return float(style_size.pt)
        except Exception:  # noqa: BLE001
            return None
        return None
    return sum(sizes) / len(sizes)


def _align(paragraph) -> str:  # noqa: ANN001
    mapping = {
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }
    return mapping.get(paragraph.alignment, "left")


def extract_docx_blocks(path: Path) -> list[TextBlock]:
    """Extrai blocos de Word preservando pistas de formatação."""
    document = Document(str(path))
    blocks: list[TextBlock] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else "Normal"
        blocks.append(
            TextBlock(
                text=text,
                style_name=style_name,
                bold=_run_bold(paragraph),
                font_size_pt=_run_font_size(paragraph),
                align=_align(paragraph),
            )
        )

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(TextBlock(text=" | ".join(cells)))

    return blocks


def extract_docx(path: Path) -> tuple[str, list[str]]:
    """Compatibilidade: retorna texto e lista simples de strings."""
    blocks = extract_docx_blocks(path)
    texts = [b.text for b in blocks]
    return "\n\n".join(texts), texts
