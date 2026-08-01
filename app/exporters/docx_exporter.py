from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.layout import LayoutSettings
from app.models import Book, Chapter


def _set_run_font(run, name: str, size: int, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _apply_book_page(document: Document, settings: LayoutSettings) -> None:
    fmt = settings.format()
    top, bottom, left, right = settings.margins_cm()
    for section in document.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)
        section.page_width = Cm(fmt.width_cm)
        section.page_height = Cm(fmt.height_cm)


def _write_chapter_body(document: Document, chapter: Chapter, settings: LayoutSettings) -> None:
    font_name = settings.font().docx_name
    size = settings.font_size
    literary = settings.style_id == "prosa_literaria"
    top_space = Pt(72 if literary else 24)

    if chapter.has_heading:
        if chapter.kind == "prologue":
            label = "Prólogo"
        elif chapter.kind == "epilogue":
            label = "Epílogo"
        elif chapter.number is not None and chapter.kind == "chapter":
            label = f"Capítulo {chapter.number}"
        else:
            label = ""

        title = chapter.title.strip()
        if chapter.kind in {"prologue", "epilogue"} and title.lower() in {
            "prólogo",
            "prologo",
            "epílogo",
            "epilogo",
        }:
            title = ""
        if label and title.lower().startswith(label.lower()):
            title = title[len(label) :].strip(" —-.")

        if label:
            cap_label = document.add_paragraph()
            cap_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_label.paragraph_format.space_before = top_space
            cap_label.paragraph_format.space_after = Pt(10)
            run = cap_label.add_run(label)
            _set_run_font(run, font_name, max(9, size - 1), italic=literary)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            heading = document.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.paragraph_format.space_after = Pt(10 if literary else 28)
        else:
            heading = document.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.paragraph_format.space_before = top_space
            heading.paragraph_format.space_after = Pt(10 if literary else 28)

        if title:
            run = heading.add_run(title)
            _set_run_font(run, font_name, size + (6 if literary else 8), bold=True)

        if settings.chapter_ornament():
            ornament = document.add_paragraph()
            ornament.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ornament.paragraph_format.space_before = Pt(6)
            ornament.paragraph_format.space_after = Pt(32)
            run = ornament.add_run("* * *")
            _set_run_font(run, font_name, size + 1)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    indent = Cm(settings.first_line_indent_cm())
    spacing = Pt(settings.paragraph_spacing_pt())

    for index, paragraph in enumerate(chapter.paragraphs):
        if paragraph.style == "section":
            if literary:
                dots = document.add_paragraph()
                dots.alignment = WD_ALIGN_PARAGRAPH.CENTER
                dots.paragraph_format.space_before = Pt(20)
                dots.paragraph_format.space_after = Pt(6)
                run = dots.add_run(". . .")
                _set_run_font(run, font_name, size)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            section = document.add_paragraph()
            section.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if literary else WD_ALIGN_PARAGRAPH.LEFT
            )
            section.paragraph_format.space_before = Pt(8 if literary else 22)
            section.paragraph_format.space_after = Pt(16)
            section.paragraph_format.first_line_indent = Cm(0)
            run = section.add_run(paragraph.text)
            _set_run_font(run, font_name, size + 1, bold=True, italic=literary)
            continue

        body = document.add_paragraph()
        body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        prev_section = index > 0 and chapter.paragraphs[index - 1].style == "section"
        if settings.skip_first_indent() and (index == 0 or prev_section):
            body.paragraph_format.first_line_indent = Cm(0)
        else:
            body.paragraph_format.first_line_indent = indent
        body.paragraph_format.space_after = spacing
        body.paragraph_format.space_before = Pt(0)
        body.paragraph_format.line_spacing = settings.line_height()
        body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        run = body.add_run(paragraph.text)
        _set_run_font(run, font_name, size)

def export_chapter_docx(
    book: Book,
    output_path: str | Path,
    settings: LayoutSettings | None = None,
) -> Path:
    settings = settings or LayoutSettings()
    path = Path(output_path)
    chapter = book.primary_chapter
    if chapter is None:
        raise ValueError("Nenhum conteúdo de capítulo para exportar.")

    document = Document()
    _apply_book_page(document, settings)
    _write_chapter_body(document, chapter, settings)
    document.save(str(path))
    return path


def export_docx(
    book: Book,
    output_path: str | Path,
    settings: LayoutSettings | None = None,
) -> Path:
    settings = settings or LayoutSettings()
    if book.is_chapter:
        return export_chapter_docx(book, output_path, settings)

    path = Path(output_path)
    font_name = settings.font().docx_name
    size = settings.font_size
    literary = settings.style_id == "prosa_literaria"
    document = Document()
    _apply_book_page(document, settings)

    for _ in range(5 if literary else 4):
        document.add_paragraph()

    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(book.title)
    _set_run_font(run, font_name, size + (14 if literary else 16), bold=True)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    if literary:
        rule = document.add_paragraph()
        rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rule.paragraph_format.space_before = Pt(16)
        run = rule.add_run("* * *")
        _set_run_font(run, font_name, size)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    if book.author:
        document.add_paragraph()
        author_para = document.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_para.add_run(book.author)
        _set_run_font(run, font_name, size + 1, italic=literary)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    document.add_page_break()

    has_named = any(
        c.has_heading or any(p.style == "section" for p in c.paragraphs)
        for c in book.chapters
    )
    if settings.include_toc and has_named:
        toc_heading = document.add_paragraph()
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = toc_heading.add_run("Sumário")
        _set_run_font(run, font_name, size + 5, bold=True)
        document.add_paragraph()

        for chapter in book.chapters:
            if chapter.has_heading:
                line = document.add_paragraph()
                line.paragraph_format.space_after = Pt(10)
                run = line.add_run(chapter.display_label)
                _set_run_font(run, font_name, size)
            for para in chapter.paragraphs:
                if para.style != "section":
                    continue
                sub = document.add_paragraph()
                sub.paragraph_format.space_after = Pt(6)
                sub.paragraph_format.left_indent = Cm(0.5)
                run = sub.add_run(para.text)
                _set_run_font(run, font_name, max(9, size - 1), italic=True)

        document.add_page_break()

    for index, chapter in enumerate(book.chapters):
        if index > 0:
            document.add_page_break()
        _write_chapter_body(document, chapter, settings)

    document.save(str(path))
    return path
